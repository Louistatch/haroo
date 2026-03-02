"""
Tests pour le moteur de templates dynamiques
"""
import os
from datetime import datetime
from io import BytesIO
from pathlib import Path

from django.test import TestCase
from openpyxl import Workbook
from docx import Document

from apps.documents.services.template_engine import (
    TemplateEngine,
    TemplateEngineError,
    VariableNotFoundError,
    UnsupportedFormatError
)


class TemplateEngineTestCase(TestCase):
    """
    Tests pour le moteur de templates dynamiques
    
    Exigences testées: 6.1, 6.2, 6.3, 6.4, 6.5, 6.6
    """
    
    def setUp(self):
        """Prépare les tests"""
        self.engine = TemplateEngine()
        self.test_dir = Path('/tmp/test_templates')
        self.test_dir.mkdir(exist_ok=True)
        
        # Créer des templates de test
        self.excel_template_path = self.test_dir / 'template.xlsx'
        self.word_template_path = self.test_dir / 'template.docx'
        
        self._create_excel_template()
        self._create_word_template()
    
    def tearDown(self):
        """Nettoie après les tests"""
        # Supprimer les fichiers de test
        import time
        import gc
        
        # Force garbage collection to close any open file handles
        gc.collect()
        time.sleep(0.1)
        
        for file in self.test_dir.glob('*'):
            try:
                file.unlink()
            except PermissionError:
                # On Windows, files might still be locked
                pass
        
        try:
            self.test_dir.rmdir()
        except OSError:
            # Directory might not be empty on Windows
            pass
    
    def _create_excel_template(self):
        """Crée un template Excel de test"""
        wb = Workbook()
        ws = wb.active
        
        # Ajouter des données avec variables
        ws['A1'] = 'Compte d\'Exploitation Prévisionnel'
        ws['A2'] = 'Région: {{region}}'
        ws['A3'] = 'Préfecture: {{prefecture}}'
        ws['A4'] = 'Canton: {{canton}}'
        ws['A5'] = 'Culture: {{culture}}'
        ws['A6'] = 'Date: {{date}}'
        ws['A7'] = 'Prix: {{prix}} FCFA'
        
        # Ajouter un tableau avec variables
        ws['A9'] = 'Description'
        ws['B9'] = 'Valeur'
        ws['A10'] = 'Localisation'
        ws['B10'] = '{{canton}}, {{prefecture}}'
        
        wb.save(str(self.excel_template_path))
        wb.close()
    
    def _create_word_template(self):
        """Crée un template Word de test"""
        doc = Document()
        
        # Ajouter du contenu avec variables
        doc.add_heading('Itinéraire Technique', 0)
        doc.add_paragraph('Région: {{region}}')
        doc.add_paragraph('Préfecture: {{prefecture}}')
        doc.add_paragraph('Canton: {{canton}}')
        doc.add_paragraph('Culture: {{culture}}')
        doc.add_paragraph('Date: {{date}}')
        doc.add_paragraph('Prix: {{prix}} FCFA')
        
        # Ajouter un tableau
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Localisation'
        table.cell(0, 1).text = '{{canton}}, {{prefecture}}'
        table.cell(1, 0).text = 'Culture'
        table.cell(1, 1).text = '{{culture}}'
        
        doc.save(str(self.word_template_path))
    
    def test_extract_variables_excel(self):
        """
        Test: Extraction des variables d'un template Excel
        Exigence: 6.2
        """
        variables = self.engine.extract_variables(
            str(self.excel_template_path),
            'EXCEL'
        )
        
        expected_variables = {'region', 'prefecture', 'canton', 'culture', 'date', 'prix'}
        self.assertEqual(variables, expected_variables)
    
    def test_extract_variables_word(self):
        """
        Test: Extraction des variables d'un template Word
        Exigence: 6.2
        """
        variables = self.engine.extract_variables(
            str(self.word_template_path),
            'WORD'
        )
        
        expected_variables = {'region', 'prefecture', 'canton', 'culture', 'date', 'prix'}
        self.assertEqual(variables, expected_variables)
    
    def test_extract_variables_unsupported_format(self):
        """
        Test: Erreur pour format non supporté
        Exigence: 6.1
        """
        with self.assertRaises(UnsupportedFormatError) as context:
            self.engine.extract_variables(
                str(self.excel_template_path),
                'PDF'
            )
        
        self.assertIn('Format non supporté', str(context.exception))
    
    def test_validate_variables_success(self):
        """
        Test: Validation réussie des variables
        Exigence: 6.3
        """
        required = ['canton', 'prefecture', 'region']
        provided = {
            'canton': 'Lomé',
            'prefecture': 'Golfe',
            'region': 'Maritime',
            'culture': 'Maïs'
        }
        
        # Ne devrait pas lever d'exception
        self.engine.validate_variables(required, provided)
    
    def test_validate_variables_missing(self):
        """
        Test: Erreur si variables requises manquantes
        Exigence: 6.3
        """
        required = ['canton', 'prefecture', 'region', 'culture']
        provided = {
            'canton': 'Lomé',
            'prefecture': 'Golfe'
        }
        
        with self.assertRaises(VariableNotFoundError) as context:
            self.engine.validate_variables(required, provided)
        
        error_message = str(context.exception)
        self.assertIn('culture', error_message)
        self.assertIn('region', error_message)
    
    def test_substitute_variables_excel(self):
        """
        Test: Substitution des variables dans Excel
        Exigences: 6.4, 6.5
        """
        variables = {
            'region': 'Maritime',
            'prefecture': 'Golfe',
            'canton': 'Lomé',
            'culture': 'Maïs',
            'date': datetime(2024, 1, 15),
            'prix': 5000
        }
        
        output = self.engine.substitute_variables(
            str(self.excel_template_path),
            'EXCEL',
            variables
        )
        
        # Vérifier que le fichier est généré
        self.assertIsInstance(output, BytesIO)
        self.assertGreater(output.getbuffer().nbytes, 0)
        
        # Charger et vérifier le contenu
        from openpyxl import load_workbook
        output.seek(0)
        wb = load_workbook(output)
        ws = wb.active
        
        # Vérifier les substitutions
        self.assertEqual(ws['A2'].value, 'Région: Maritime')
        self.assertEqual(ws['A3'].value, 'Préfecture: Golfe')
        self.assertEqual(ws['A4'].value, 'Canton: Lomé')
        self.assertEqual(ws['A5'].value, 'Culture: Maïs')
        self.assertEqual(ws['A6'].value, 'Date: 15/01/2024')
        self.assertEqual(ws['A7'].value, 'Prix: 5 000 FCFA')
        self.assertEqual(ws['B10'].value, 'Lomé, Golfe')
        
        # Vérifier qu'aucune variable n'est restée
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    self.assertNotIn('{{', cell.value)
                    self.assertNotIn('}}', cell.value)
        
        wb.close()
    
    def test_substitute_variables_word(self):
        """
        Test: Substitution des variables dans Word
        Exigences: 6.4, 6.5
        """
        variables = {
            'region': 'Maritime',
            'prefecture': 'Golfe',
            'canton': 'Lomé',
            'culture': 'Maïs',
            'date': datetime(2024, 1, 15),
            'prix': 5000
        }
        
        output = self.engine.substitute_variables(
            str(self.word_template_path),
            'WORD',
            variables
        )
        
        # Vérifier que le fichier est généré
        self.assertIsInstance(output, BytesIO)
        self.assertGreater(output.getbuffer().nbytes, 0)
        
        # Charger et vérifier le contenu
        output.seek(0)
        doc = Document(output)
        
        # Récupérer tout le texte
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Vérifier les substitutions
        self.assertIn('Région: Maritime', full_text)
        self.assertIn('Préfecture: Golfe', full_text)
        self.assertIn('Canton: Lomé', full_text)
        self.assertIn('Culture: Maïs', full_text)
        self.assertIn('Date: 15/01/2024', full_text)
        self.assertIn('Prix: 5 000 FCFA', full_text)
        
        # Vérifier qu'aucune variable n'est restée
        self.assertNotIn('{{', full_text)
        self.assertNotIn('}}', full_text)
        
        # Vérifier les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.assertNotIn('{{', cell.text)
                    self.assertNotIn('}}', cell.text)
    
    def test_normalize_variables_date(self):
        """
        Test: Normalisation des dates au format français
        Exigence: 6.5
        """
        variables = {
            'date': datetime(2024, 3, 25)
        }
        
        normalized = self.engine._normalize_variables(variables)
        self.assertEqual(normalized['date'], '25/03/2024')
    
    def test_normalize_variables_numbers(self):
        """
        Test: Normalisation des nombres avec séparateurs
        Exigence: 6.5
        """
        variables = {
            'prix': 5000,
            'superficie': 12.5
        }
        
        normalized = self.engine._normalize_variables(variables)
        self.assertEqual(normalized['prix'], '5 000')
        self.assertEqual(normalized['superficie'], '12.50')
    
    def test_normalize_variables_none(self):
        """
        Test: Normalisation des valeurs None
        Exigence: 6.5
        """
        variables = {
            'optional': None
        }
        
        normalized = self.engine._normalize_variables(variables)
        self.assertEqual(normalized['optional'], '')
    
    def test_substitute_with_missing_variables(self):
        """
        Test: Les variables non fournies restent inchangées
        Exigence: 6.5
        """
        variables = {
            'region': 'Maritime',
            'canton': 'Lomé'
        }
        
        output = self.engine.substitute_variables(
            str(self.excel_template_path),
            'EXCEL',
            variables
        )
        
        output.seek(0)
        from openpyxl import load_workbook
        wb = load_workbook(output)
        ws = wb.active
        
        # Les variables fournies doivent être substituées
        self.assertEqual(ws['A2'].value, 'Région: Maritime')
        self.assertEqual(ws['A4'].value, 'Canton: Lomé')
        
        # Les variables non fournies doivent rester
        self.assertEqual(ws['A3'].value, 'Préfecture: {{prefecture}}')
        self.assertEqual(ws['A5'].value, 'Culture: {{culture}}')
        
        wb.close()
    
    def test_excel_template_with_formulas(self):
        """
        Test: Les formules Excel sont préservées
        Exigence: 6.5
        """
        # Créer un template avec formules
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Prix unitaire: {{prix}}'
        ws['A2'] = 'Quantité: 10'
        ws['A3'] = '=A1*A2'  # Formule
        
        formula_path = self.test_dir / 'formula_template.xlsx'
        wb.save(str(formula_path))
        wb.close()
        
        variables = {'prix': 100}
        output = self.engine.substitute_variables(
            str(formula_path),
            'EXCEL',
            variables
        )
        
        output.seek(0)
        from openpyxl import load_workbook
        wb = load_workbook(output)
        ws = wb.active
        
        # La formule doit être préservée
        self.assertTrue(ws['A3'].value.startswith('='))
        
        wb.close()
    
    def test_word_template_preserves_formatting(self):
        """
        Test: Le formatage Word est préservé
        Exigence: 6.5
        """
        # Créer un template avec formatage
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run('Canton: {{canton}}')
        run.bold = True
        run.italic = True
        
        format_path = self.test_dir / 'format_template.docx'
        doc.save(str(format_path))
        
        variables = {'canton': 'Lomé'}
        output = self.engine.substitute_variables(
            str(format_path),
            'WORD',
            variables
        )
        
        output.seek(0)
        doc = Document(output)
        
        # Le texte doit être substitué
        self.assertEqual(doc.paragraphs[0].text, 'Canton: Lomé')
        
        # Le formatage doit être préservé
        run = doc.paragraphs[0].runs[0]
        self.assertTrue(run.bold)
        self.assertTrue(run.italic)
    
    def test_multiple_variables_in_same_cell(self):
        """
        Test: Plusieurs variables dans la même cellule
        Exigence: 6.5
        """
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '{{canton}}, {{prefecture}}, {{region}}'
        
        multi_path = self.test_dir / 'multi_template.xlsx'
        wb.save(str(multi_path))
        wb.close()
        
        variables = {
            'canton': 'Lomé',
            'prefecture': 'Golfe',
            'region': 'Maritime'
        }
        
        output = self.engine.substitute_variables(
            str(multi_path),
            'EXCEL',
            variables
        )
        
        output.seek(0)
        from openpyxl import load_workbook
        wb = load_workbook(output)
        ws = wb.active
        
        self.assertEqual(ws['A1'].value, 'Lomé, Golfe, Maritime')
        
        wb.close()
    
    def test_case_sensitive_variables(self):
        """
        Test: Les variables sont sensibles à la casse
        Exigence: 6.2
        """
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '{{Canton}}'  # Majuscule
        ws['A2'] = '{{canton}}'  # Minuscule
        
        case_path = self.test_dir / 'case_template.xlsx'
        wb.save(str(case_path))
        wb.close()
        
        # Extraire les variables
        variables = self.engine.extract_variables(str(case_path), 'EXCEL')
        
        # Les deux doivent être détectées comme différentes
        self.assertIn('Canton', variables)
        self.assertIn('canton', variables)
        self.assertEqual(len(variables), 2)
        
        wb.close()
