"""
Moteur de templates dynamiques pour la génération de documents personnalisés
Supporte les formats Excel (.xlsx) et Word (.docx)
"""
import re
from typing import Dict, List, Set, Optional, Any
from pathlib import Path
from datetime import datetime
from io import BytesIO

from openpyxl import load_workbook
from openpyxl.workbook import Workbook
from docx import Document
from django.core.files.base import ContentFile


class TemplateEngineError(Exception):
    """Exception de base pour les erreurs du moteur de templates"""
    pass


class VariableNotFoundError(TemplateEngineError):
    """Exception levée quand une variable requise n'est pas fournie"""
    pass


class UnsupportedFormatError(TemplateEngineError):
    """Exception levée pour un format de fichier non supporté"""
    pass


class TemplateEngine:
    """
    Moteur de templates dynamiques pour substitution de variables
    
    Formats supportés:
    - Excel (.xlsx)
    - Word (.docx)
    
    Variables supportées:
    - {{canton}}: Nom du canton
    - {{prefecture}}: Nom de la préfecture
    - {{region}}: Nom de la région
    - {{culture}}: Type de culture
    - {{date}}: Date actuelle ou spécifiée
    - {{prix}}: Prix du document ou autre valeur monétaire
    
    Exigences: 6.1, 6.2, 6.3, 6.4, 6.5, 6.6
    """
    
    # Pattern pour identifier les variables: {{nom_variable}}
    VARIABLE_PATTERN = re.compile(r'\{\{(\w+)\}\}')
    
    # Variables supportées par la plateforme
    SUPPORTED_VARIABLES = {
        'canton', 'prefecture', 'region', 'culture', 'date', 'prix'
    }
    
    def __init__(self):
        """Initialise le moteur de templates"""
        pass
    
    def extract_variables(self, file_path: str, file_format: str) -> Set[str]:
        """
        Extrait toutes les variables présentes dans un template
        
        Args:
            file_path: Chemin vers le fichier template
            file_format: Format du fichier ('EXCEL' ou 'WORD')
            
        Returns:
            Set des noms de variables trouvées
            
        Raises:
            UnsupportedFormatError: Si le format n'est pas supporté
            
        Exigences: 6.2
        """
        if file_format == 'EXCEL':
            return self._extract_variables_excel(file_path)
        elif file_format == 'WORD':
            return self._extract_variables_word(file_path)
        else:
            raise UnsupportedFormatError(
                f"Format non supporté: {file_format}. "
                f"Formats supportés: EXCEL, WORD"
            )
    
    def _extract_variables_excel(self, file_path: str) -> Set[str]:
        """
        Extrait les variables d'un fichier Excel
        
        Args:
            file_path: Chemin vers le fichier Excel
            
        Returns:
            Set des noms de variables trouvées
        """
        variables = set()
        workbook = load_workbook(file_path)
        
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        matches = self.VARIABLE_PATTERN.findall(cell.value)
                        variables.update(matches)
        
        workbook.close()
        return variables
    
    def _extract_variables_word(self, file_path: str) -> Set[str]:
        """
        Extrait les variables d'un fichier Word
        
        Args:
            file_path: Chemin vers le fichier Word
            
        Returns:
            Set des noms de variables trouvées
        """
        variables = set()
        document = Document(file_path)
        
        # Extraire des paragraphes
        for paragraph in document.paragraphs:
            if paragraph.text:
                matches = self.VARIABLE_PATTERN.findall(paragraph.text)
                variables.update(matches)
        
        # Extraire des tableaux
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        matches = self.VARIABLE_PATTERN.findall(cell.text)
                        variables.update(matches)
        
        # Extraire des en-têtes et pieds de page
        for section in document.sections:
            # En-tête
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text:
                        matches = self.VARIABLE_PATTERN.findall(paragraph.text)
                        variables.update(matches)
            
            # Pied de page
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text:
                        matches = self.VARIABLE_PATTERN.findall(paragraph.text)
                        variables.update(matches)
        
        return variables
    
    def validate_variables(
        self,
        required_variables: List[str],
        provided_variables: Dict[str, Any]
    ) -> None:
        """
        Valide que toutes les variables requises sont fournies
        
        Args:
            required_variables: Liste des variables requises
            provided_variables: Dictionnaire des variables fournies
            
        Raises:
            VariableNotFoundError: Si des variables requises manquent
            
        Exigences: 6.3
        """
        missing_variables = set(required_variables) - set(provided_variables.keys())
        
        if missing_variables:
            raise VariableNotFoundError(
                f"Variables requises manquantes: {', '.join(sorted(missing_variables))}"
            )
    
    def substitute_variables(
        self,
        template_path: str,
        file_format: str,
        variables: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> BytesIO:
        """
        Substitue les variables dans un template et génère le document final
        
        Args:
            template_path: Chemin vers le fichier template
            file_format: Format du fichier ('EXCEL' ou 'WORD')
            variables: Dictionnaire des variables et leurs valeurs
            output_path: Chemin optionnel pour sauvegarder le fichier
            
        Returns:
            BytesIO contenant le document généré
            
        Raises:
            UnsupportedFormatError: Si le format n'est pas supporté
            
        Exigences: 6.4, 6.5
        """
        # Normaliser les valeurs des variables
        normalized_vars = self._normalize_variables(variables)
        
        if file_format == 'EXCEL':
            return self._substitute_variables_excel(
                template_path, normalized_vars, output_path
            )
        elif file_format == 'WORD':
            return self._substitute_variables_word(
                template_path, normalized_vars, output_path
            )
        else:
            raise UnsupportedFormatError(
                f"Format non supporté: {file_format}. "
                f"Formats supportés: EXCEL, WORD"
            )
    
    def _normalize_variables(self, variables: Dict[str, Any]) -> Dict[str, str]:
        """
        Normalise les valeurs des variables en chaînes de caractères
        
        Args:
            variables: Dictionnaire des variables
            
        Returns:
            Dictionnaire avec valeurs normalisées
        """
        normalized = {}
        
        for key, value in variables.items():
            if value is None:
                normalized[key] = ''
            elif isinstance(value, datetime):
                # Format de date français: JJ/MM/AAAA
                normalized[key] = value.strftime('%d/%m/%Y')
            elif isinstance(value, (int, float)):
                # Formater les nombres avec séparateur de milliers
                if isinstance(value, float):
                    normalized[key] = f"{value:,.2f}".replace(',', ' ')
                else:
                    normalized[key] = f"{value:,}".replace(',', ' ')
            else:
                normalized[key] = str(value)
        
        return normalized
    
    def _substitute_variables_excel(
        self,
        template_path: str,
        variables: Dict[str, str],
        output_path: Optional[str] = None
    ) -> BytesIO:
        """
        Substitue les variables dans un fichier Excel
        
        Args:
            template_path: Chemin vers le template Excel
            variables: Dictionnaire des variables normalisées
            output_path: Chemin optionnel pour sauvegarder
            
        Returns:
            BytesIO contenant le document généré
        """
        workbook = load_workbook(template_path)
        
        # Parcourir toutes les feuilles
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        # Remplacer toutes les variables dans la cellule
                        new_value = cell.value
                        for var_name, var_value in variables.items():
                            pattern = f'{{{{{var_name}}}}}'
                            new_value = new_value.replace(pattern, var_value)
                        cell.value = new_value
        
        # Sauvegarder dans un BytesIO
        output = BytesIO()
        workbook.save(output)
        output.seek(0)
        
        # Sauvegarder sur disque si demandé
        if output_path:
            workbook.save(output_path)
        
        workbook.close()
        return output
    
    def _substitute_variables_word(
        self,
        template_path: str,
        variables: Dict[str, str],
        output_path: Optional[str] = None
    ) -> BytesIO:
        """
        Substitue les variables dans un fichier Word
        
        Args:
            template_path: Chemin vers le template Word
            variables: Dictionnaire des variables normalisées
            output_path: Chemin optionnel pour sauvegarder
            
        Returns:
            BytesIO contenant le document généré
        """
        document = Document(template_path)
        
        # Fonction helper pour remplacer dans un texte
        def replace_variables(text: str) -> str:
            for var_name, var_value in variables.items():
                pattern = f'{{{{{var_name}}}}}'
                text = text.replace(pattern, var_value)
            return text
        
        # Remplacer dans les paragraphes
        for paragraph in document.paragraphs:
            if paragraph.text:
                # Préserver le formatage en remplaçant run par run
                for run in paragraph.runs:
                    if run.text:
                        run.text = replace_variables(run.text)
        
        # Remplacer dans les tableaux
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = replace_variables(run.text)
        
        # Remplacer dans les en-têtes et pieds de page
        for section in document.sections:
            # En-tête
            if section.header:
                for paragraph in section.header.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = replace_variables(run.text)
            
            # Pied de page
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = replace_variables(run.text)
        
        # Sauvegarder dans un BytesIO
        output = BytesIO()
        document.save(output)
        output.seek(0)
        
        # Sauvegarder sur disque si demandé
        if output_path:
            document.save(output_path)
        
        return output
    
    def generate_document(
        self,
        template_file,
        file_format: str,
        variables: Dict[str, Any],
        required_variables: Optional[List[str]] = None
    ) -> ContentFile:
        """
        Génère un document à partir d'un template avec validation
        
        Args:
            template_file: Fichier template (Django File object)
            file_format: Format du fichier ('EXCEL' ou 'WORD')
            variables: Dictionnaire des variables et leurs valeurs
            required_variables: Liste optionnelle des variables requises
            
        Returns:
            ContentFile prêt à être sauvegardé dans un FileField
            
        Raises:
            VariableNotFoundError: Si des variables requises manquent
            UnsupportedFormatError: Si le format n'est pas supporté
            
        Exigences: 6.3, 6.4, 6.5
        """
        # Valider les variables si une liste est fournie
        if required_variables:
            self.validate_variables(required_variables, variables)
        
        # Sauvegarder temporairement le template
        temp_path = f'/tmp/template_{datetime.now().timestamp()}'
        with open(temp_path, 'wb') as f:
            for chunk in template_file.chunks():
                f.write(chunk)
        
        try:
            # Générer le document
            output = self.substitute_variables(
                temp_path,
                file_format,
                variables
            )
            
            # Créer un ContentFile pour Django
            extension = 'xlsx' if file_format == 'EXCEL' else 'docx'
            filename = f'document_{datetime.now().strftime("%Y%m%d_%H%M%S")}.{extension}'
            
            return ContentFile(output.getvalue(), name=filename)
        
        finally:
            # Nettoyer le fichier temporaire
            Path(temp_path).unlink(missing_ok=True)
